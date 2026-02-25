import re
from typing import Dict
from docx import Document

def replace_placeholders_in_doc(doc: Document, variables: Dict[str, str]) -> None:
    """
    Заменяет все плейсхолдеры вида {{variable_name}} в документе docx
    на соответствующие значения из словаря variables.
    Также удаляет аннотации вида (string; user_input) и подобные.

    Аргументы:
        doc: загруженный документ python-docx
        variables: словарь {имя_переменной: значение_строкой}
    """
    placeholder_pattern = re.compile(r"\{\{([^}]+)\}\}")
    annotation_pattern = re.compile(r"\s+\([^;)]+;[^)]+\)")

    def replace_in_text(text: str) -> str:
        if not text:
            return text
        # Замена плейсхолдеров
        text = placeholder_pattern.sub(
            lambda m: variables.get(m.group(1).strip(), m.group(0)), text
        )
        # Удаление аннотаций
        text = annotation_pattern.sub("", text)
        return text

    # Обработка всех параграфов
    for paragraph in doc.paragraphs:
        if paragraph.text:
            new_text = replace_in_text(paragraph.text)
            if new_text != paragraph.text:
                paragraph.clear()
                paragraph.add_run(new_text)

    # Обработка всех ячеек таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text:
                        new_text = replace_in_text(paragraph.text)
                        if new_text != paragraph.text:
                            paragraph.clear()
                            paragraph.add_run(new_text)